# my_streamlit_app.py
import streamlit as st
from docx import Document
from pydub import AudioSegment
import speech_recognition as sr
import tempfile
import os

# Initialize the recognizer
recognizer = sr.Recognizer()

# Function to transcribe audio and save as DOCX file
def transcribe_audio(audio_file, docx_file):
    with sr.AudioFile(audio_file) as source:
        audio = recognizer.record(source)  # Record the entire audio file
        try:
            text = recognizer.recognize_google(audio)  # Use Google Web Speech API
            doc = Document()
            doc.add_heading("Transcription", level=1)
            doc.add_paragraph(text)
            doc.save(docx_file)
            return True
        except sr.UnknownValueError:
            return False
        except sr.RequestError as e:
            st.error(f"Could not request results from Google Web Speech API; {e}")
            return False

def main():
    st.title("Audio to Text Transcription App")

    # Upload WAV file
    uploaded_file = st.file_uploader("Upload a WAV file", type=["wav"])

    if uploaded_file:
        st.audio(uploaded_file, format="audio/wav")

    # Convert button
    if st.button("Convert"):
        if uploaded_file:
            # Create a temporary directory to store the DOCX file
            with tempfile.TemporaryDirectory() as temp_dir:
                docx_file = os.path.join(temp_dir, "transcription.docx")

                # Convert and save the transcription as a DOCX file
                if transcribe_audio(uploaded_file, docx_file):
                    st.success("Transcription completed.")

                    # Download button for the DOCX file
                    st.download_button(
                        label="Download DOCX",
                        data=open(docx_file, "rb").read(),
                        file_name="transcription.docx",
                        key="docx_download",
                    )
                else:
                    st.error("Transcription failed. Please check the audio file and try again.")
        else:
            st.error("Please upload a WAV audio file.")

if __name__ == "__main__":
    main()
