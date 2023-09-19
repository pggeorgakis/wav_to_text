import speech_recognition as sr
from docx import Document

# Initialize the recognizer
recognizer = sr.Recognizer()

# Load the MP3 audio file
audio_file = r"C:\Users\pgeor\Documents\code\misc\mp3_to_text\test.wav"  # Replace with the path to your MP3 file

# Create a Word document
doc = Document()
doc.add_heading("Transcription", level=1)

# Function to transcribe audio and add to Word document
def transcribe_audio(audio_file):
    with sr.AudioFile(audio_file) as source:
        audio = recognizer.record(source)  # Record the entire audio file
        try:
            text = recognizer.recognize_google(audio)  # Use Google Web Speech API
            doc.add_paragraph(text)  # Add transcribed text to the Word document
        except sr.UnknownValueError:
            print("Google Web Speech API could not understand audio.")
        except sr.RequestError as e:
            print(f"Could not request results from Google Web Speech API; {e}")

# Transcribe the audio and add to the Word document
transcribe_audio(audio_file)

# Save the Word document
doc.save("transcription.docx")

print("Transcription saved to transcription.docx")

